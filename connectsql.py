

import pandas as pd
import jieba

import string
import time
import traceback

import pymysql
import requests
import re

import xlwt
from bs4 import BeautifulSoup
from flask import json
import nltk
from nltk.corpus import brown

import fileroot

#数据库的连接查询
def get_conn_mysql():
    """
    :return: 连接，游标192.168.1.102
    """
    # 创建连接
    conn = pymysql.connect(host="127.0.0.1",
                    user="root",
                    password="123456",
                    db="bigwork_data",
                    charset="utf8")
    # 创建游标
    cursor = conn.cursor()  # 执行完毕返回的结果集默认以元组显示
    return conn, cursor
def get_conn_mysql_name(name):
    """
    :return: 连接，游标192.168.1.102
    """
    # 创建连接
    conn = pymysql.connect(host="127.0.0.1",
                    user="root",
                    password="123456",
                    db=name,
                    charset="utf8")
    # 创建游标
    cursor = conn.cursor()  # 执行完毕返回的结果集默认以元组显示
    return conn, cursor
def close_conn_mysql(conn, cursor):
    if cursor:
        cursor.close()
    if conn:
        conn.close()
def query_mysql(sql,*args):
    """
    封装通用查询
    :param sql:
    :param args:
    :return: 返回查询结果以((),(),)形式
    """
    conn,cursor = get_conn_mysql();
    cursor.execute(sql)
    res=cursor.fetchall()
    close_conn_mysql(conn,cursor)
    return res


#读取样表生成数据字典type_flag 1 追加，2 覆盖
def read_example(path,china_row,english_row,unit_row,type_flag):
    flag=1
    conn, cursor = get_conn_mysql()
    #将excel转换为csv文件
    data = pd.read_excel('import_excel/'+path, 'Sheet1')
    data.fillna('', inplace=True)
    print(data)
    csv_name = path.split(".")[0]
    # data.to_csv("excel_data/"+csv_name+'.csv', encoding='utf-8')
    # data_csv=pd.read_csv("excel_data/"+csv_name+".csv")
    # 编写表创建语句(字段类型就设为string)
    # 表名
    table_name = path.split(".")[0]
    ######################################################
    #获取数据三种情况各自为keys
    if(china_row=="1"):
        key_china = data.keys()
        key_0 = data.values.tolist()[(int)(english_row)-2]
        key_unit=data.values.tolist()[(int)(unit_row)-2]
    if(english_row=="1"):
        key_china = data.values.tolist()[(int)(china_row)-2]
        key_0 = data.keys()
        key_unit=data.values.tolist()[(int)(unit_row)-2]
    if(unit_row=="1"):
        key_china = data.values.tolist()[(int)(china_row)-2]
        key_0 = data.values.tolist()[(int)(english_row)-2]
        key_unit=data.keys()
    key=""
    for i in key_0:
        key=key+","+i
    key=key[1:]
    if (type_flag=="0"):
        delete_sql = "drop table " + csv_name
        try:
            cursor.execute(delete_sql)
        except:
            traceback.print_exc()
            flag = 0
            print("表删除失败")
    #建表及插入数据
    sql = "CREATE TABLE IF NOT EXISTS " + csv_name + " ("
    # 循环加入key值
    j=0
    for i in key_0:
        sql = sql + i + " TEXT  comment '"+key_china[j]+","+key_unit[j]+"',"
        j=j+1;
    creat_sql = sql[0:-1] + ") ENGINE = InnoDB DEFAULT CHARACTER SET = utf8 COLLATE = utf8_bin;"
    print(creat_sql)
    # 获取%s
    s = ','.join(['%s' for _ in range(len(data.columns))])
    # 获取values
    values=[]
    for i in data.values.tolist()[2:]:
        values.append(i)
    # 组装insert语句
    insert_sql = 'insert into {}({}) values({})'.format(table_name, key, s)
    print(insert_sql)
    # print(key_china)#中文字段名
    # print(key_unit)#中文单位
    # print(key)#英文字段名
    # print(values)#数据
    #创建表
    try:
        cursor.execute(creat_sql)
    except:
        traceback.print_exc()
        flag=0
        print("表创建失败")
    # 插入数据
    try:
        for i in values:
            cursor.execute(insert_sql, i)
            print(insert_sql)
            print(i)
        conn.commit()
    except:
        traceback.print_exc()
        flag=0
        print("写入错误")
    close_conn_mysql(cursor, conn)
    return flag
# 读取csv文件
def read_csv(path,china_row,english_row,unit_row,type_flag):
    conn, cursor=get_conn_mysql()
    flag=1
    data=pd.read_csv("import_csv/"+path)
    data.fillna('', inplace=True)
    #编写表创建语句(字段类型就设为string)
    #表名
    table_name = path.split(".")[0]
    #########################################################
    # 获取数据
    print(china_row,english_row,unit_row)
    if (china_row == "1"):
        key_china = data.keys()
        key_0 = data.values.tolist()[(int)(english_row) - 2]
        key_unit = data.values.tolist()[(int)(unit_row) - 2]
    if (english_row == "1"):
        key_china = data.values.tolist()[(int)(china_row) - 2]
        key_0 = data.keys()
        key_unit = data.values.tolist()[(int)(unit_row) - 2]
    if (unit_row == "1"):
        key_china = data.values.tolist()[(int)(china_row) - 2]
        key_0 = data.values.tolist()[(int)(english_row) - 2]
        key_unit = data.keys()
    key=""
    for i in key_0:
        key=key+","+i
    key=key[1:]
    if (type_flag == "0"):
        if(fileroot.find_filedata_filename(table_name)!=()):
            delete_sql = "drop table " + table_name
            try:
                cursor.execute(delete_sql)
            except:
                traceback.print_exc()
                flag = 0
                print("表删除失败")
    # 建表及插入数据
    sql = "CREATE TABLE IF NOT EXISTS " + table_name + " ("
    # 循环加入key值
    j = 0
    for i in key_0:
        sql = sql + i + " TEXT  comment '" + key_china[j]+","+key_unit[j]+ "',"
        j = j + 1;
    creat_sql = sql[0:-1] + ") ENGINE = InnoDB DEFAULT CHARACTER SET = utf8 COLLATE = utf8_bin;"
    print(creat_sql)
    # 获取%s
    s = ','.join(['%s' for _ in range(len(data.columns))])
    # 获取values
    values = []
    for i in data.values.tolist()[2:]:
        values.append(i)
    # 组装insert语句
    insert_sql = 'insert into {}({}) values({})'.format(table_name, key, s)
    print(insert_sql)
    # print(insert_sql)
    # print(key_china)中文字段名
    # print(key_unit)中文单位
    # print(key)英文字段名
    # print(values)数据
    # 创建表
    try:
        cursor.execute(creat_sql)
    except:
        traceback.print_exc()
        flag = 0
        print("表创建失败")
    # # 插入数据
    try:
        for i in values:
            cursor.execute(insert_sql, i)
            print(insert_sql)
            print(i)
        conn.commit()
    except:
        traceback.print_exc()
        flag = 0
        print("写入错误")
    close_conn_mysql(cursor, conn)
    return flag
#指定列导出为excel文件
#先从数据库读出指定列，在转换为excel文件，返回文件名，在让用户下载
def export_excel(getdata_str,table_name,database_name):
    # 连接数据库，查询数据
    conn,cur=get_conn_mysql_name(database_name)
    #将最后的逗号去掉
    getdata_str=getdata_str[0:-1]
    sql = 'select '+getdata_str+' from %s' % table_name
    print(sql)
    cur.execute(sql)  # 返回受影响的行数

    fields = [field[0] for field in cur.description]  # 获取所有字段名
    all_data = cur.fetchall()  # 所有数据

    # 写入excel
    book = xlwt.Workbook()
    sheet = book.add_sheet('sheet1')

    for col, field in enumerate(fields):
        sheet.write(0, col, field)

    row = 1
    for data in all_data:
        for col, field in enumerate(data):
            sheet.write(row, col, field)
        row += 1
    book.save("export_excel_select/%s.xls" % table_name)
    print("文件已导出："+"export_excel_select/%s.xls" % table_name)
    close_conn_mysql(conn,cur)
    return "%s.xls" % table_name

# 读取docx中的文本代码示例
import docx
#word读取测试
def read_word_test():
    # 获取文档对象
    file = docx.Document("import_word/word.docx")
    print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段

    # 输出每一段的内容
    for para in file.paragraphs:
        print(para.text)

    # # 输出段落编号及段落内容
    # for i in range(len(file.paragraphs)):
    #     print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)
# 读取excel文件
def read_excel(path):
    conn, cursor = get_conn_mysql()
    #将excel转换为csv文件
    data = pd.read_excel('excel_data/'+path, 'Sheet1')
    csv_name = path.split(".")[0]
    # data.to_csv("excel_data/"+csv_name+'.csv', encoding='utf-8')
    # data_csv=pd.read_csv("excel_data/"+csv_name+".csv")
    # 编写表创建语句(字段类型就设为string)
    # 表名
    table_name = path.split(".")[0]
    sql = "CREATE TABLE " + csv_name + " ("
    # 获取key值 CREATE TABLE `bigwork_data`.`table_test` (
    #   `id` VARCHAR(45) NOT NULL,
    #   `table_testcol` VARCHAR(45) NOT NULL,
    #   `table_testcol1` VARCHAR(45) NOT NULL,
    #   `table_testcol2` VARCHAR(45) NOT NULL)
    # ENGINE = InnoDB
    # DEFAULT CHARACTER SET = utf8
    # COLLATE = utf8_bin;
    # 循环加入key值
    keys = ""
    for i in data.keys():
        sql = sql + i + " VARCHAR(45) NOT NULL,"
        keys = keys + i + ","
    keys = keys[0:-1]
    creat_sql = sql[0:-1] + ") ENGINE = InnoDB DEFAULT CHARACTER SET = utf8 COLLATE = utf8_bin;"
    # 获取%s
    s = ','.join(['%s' for _ in range(len(data.columns))])
    # 获取values
    values = data.values.tolist()
    print(values)
    # 组装insert语句
    insert_sql = 'insert into {}({}) values({})'.format(table_name, keys, s)
    print(insert_sql)
    print(creat_sql)
    print(keys);
    print(values)
    # # 创建表
    # try:
    #     cursor.execute(creat_sql)
    # except:
    #     traceback.print_exc()
    #     print("表创建失败")
    # # # 插入数据
    # try:
    #     for i in values:
    #         cursor.execute(insert_sql, i)
    #         print(insert_sql)
    #         print(i)
    #     conn.commit()
    # except:
    #     traceback.print_exc()
    #     print("写入错误")
    close_conn_mysql(cursor, conn)


if __name__ == '__main__':
    export_excel("test1,test2,", "us_video_data_numbers")
    pass
